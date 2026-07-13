import streamlit as st
import openai
import pdfplumber
import io
import re
import os
import zipfile
import json
import xml.etree.ElementTree as ET
import github
from github import Github
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- Настройка страницы ---
st.set_page_config(
    page_title="MSDS Yandex AI Studio Pro",
    page_icon="🧪",
    layout="wide"
)

# --- Автоматическое чтение конфигураций из Streamlit Secrets (по секциям) ---
yandex_secrets = st.secrets.get("yandex", {})
FOLDER_ID = yandex_secrets.get("folder_id", "")
API_KEY = yandex_secrets.get("api_key", "")

github_secrets = st.secrets.get("github", {})
GITHUB_TOKEN = github_secrets.get("token", "")
GITHUB_REPO = github_secrets.get("repo", "")
TARGET_BRANCH = github_secrets.get("branch", "main")


def clean_inline_duplicate(text: str) -> str:
    """Удаляет дублирование фраз, склеенных внутри одной строки (например, 'HelloHello')"""
    text = text.strip()
    if not text:
        return ""
    mid = len(text) // 2
    if len(text) % 2 == 0 and text[:mid] == text[mid:]:
        return text[:mid].strip()
    return text

def extract_raw_xml_text_from_zip(file_bytes) -> str:
    """Глубокое извлечение текста из DOCX, включая текстовые поля и фигуры."""
    WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    PARA_TAG = f'{WORD_NAMESPACE}p'
    TEXT_TAG = f'{WORD_NAMESPACE}t'
    NUM_PR_TAG = f'.//{WORD_NAMESPACE}numPr'

    all_extracted_lines = []
    
    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            xml_files = [f for f in z.namelist() if f.endswith('.xml') and f.startswith('word/')]
            xml_files.sort(key=lambda x: ('document' in x, x))

            for xml_file in xml_files:
                with z.open(xml_file) as f:
                    root = ET.fromstring(f.read())
                    for p in root.iter(PARA_TAG):
                        prefix = "• " if p.find(NUM_PR_TAG) is not None else ""
                        text_pieces = [node.text for node in p.iter(TEXT_TAG) if node.text]
                        p_text = "".join(text_pieces).strip()
                        p_text = clean_inline_duplicate(p_text)
                        
                        if p_text:
                            if not all_extracted_lines or all_extracted_lines[-1] != f"{prefix}{p_text}":
                                all_extracted_lines.append(f"{prefix}{p_text}")
    except Exception as e:
        st.error(f"Ошибка при XML-парсинге DOCX: {e}")
        return ""

    return "\n".join(all_extracted_lines)

def parse_line(line_str: str) -> dict:
    """
    Единый интеллектуальный парсер строки. 
    Гарантирует 100% совпадение структуры на этапах экстракции и сборки.
    """
    line_str = line_str.strip()
    if not line_str:
        return {"type": "empty"}
        
    # 1. Главные разделы (SECTION 1, РАЗДЕЛ 2)
    if bool(re.match(r'(?im)^[ \t]*(?:section|раздел)\s*\d+', line_str)):
        return {"type": "section", "text": line_str}
        
    # 2. Цифровые подразделы любого уровня (1.1, 2.3.1, 15.2.4 и т.д.)
    sub_match = re.match(r'^(\d+\.\d+(?:\.\d+)*\.?)\s*(.*)$', line_str)
    if sub_match:
        num = sub_match.group(1)
        rest = sub_match.group(2).strip()
        if ':' in rest:
            key, val = rest.split(':', 1)
            return {"type": "subsection", "num": num, "key": key.strip(), "val": val.strip()}
        else:
            return {"type": "subsection", "num": num, "key": rest, "val": None}

    # 3. Обнаружение табличной структуры (разделение табами или 3+ пробелами)
    chunks = [c.strip() for c in re.split(r'\t|\s{3,}', line_str) if c.strip()]
    if len(chunks) > 2:
        return {"type": "table_row", "chunks": chunks}
    if len(chunks) == 2 and ':' not in line_str:
        return {"type": "table_row", "chunks": chunks}
        
    # 4. Стандартные параметры формата "Ключ: Значение"
    if ':' in line_str:
        key, val = line_str.split(':', 1)
        key_str = key.strip()
        val_str = val.strip()
        # Проверяем, что это не ссылка и длина ключа адекватна
        if key_str and len(key_str) < 100 and not key_str.lower().startswith(('http', 'www')):
            return {"type": "key_value", "key": key_str, "val": val_str}
            
    # 5. Обычный текст или маркированный список
    is_bullet = line_str.startswith(('•', '-', '*'))
    clean_text = line_str.lstrip('•-* ').strip()
    return {"type": "text", "text": clean_text, "is_bullet": is_bullet}

def is_technical_garbage(s: str) -> bool:
    """Определяет, состоит ли строка только из цифр, кодов, CAS или спецсимволов."""
    s_clean = s.strip()
    if not s_clean:
        return True
    if re.match(r'^[\d\s\.,\-\/\\#№:;%()\*\+\[\]\>\<\|\=]+$', s_clean):
        return True
    return False

def extract_translation_candidates(text: str) -> set:
    """Извлекает чистые текстовые кандидаты на перевод на основе единого парсера."""
    candidates = set()
    stop_words = ['www.', 'http', 'safety data sheet', 'material safety data sheet']
    
    for line in text.split('\n'):
        line_str = line.strip()
        if not line_str or any(sw in line_str.lower() for sw in stop_words):
            continue
            
        parsed = parse_line(line_str)
        
        if parsed["type"] == "section":
            if not is_technical_garbage(parsed["text"]):
                candidates.add(parsed["text"])
                
        elif parsed["type"] == "subsection":
            if parsed["key"] and not is_technical_garbage(parsed["key"]):
                candidates.add(parsed["key"])
            if parsed["val"] and not is_technical_garbage(parsed["val"]):
                candidates.add(parsed["val"])
                
        elif parsed["type"] == "key_value":
            if parsed["key"] and not is_technical_garbage(parsed["key"]):
                candidates.add(parsed["key"])
            if parsed["val"] and not is_technical_garbage(parsed["val"]):
                candidates.add(parsed["val"])
                
        elif parsed["type"] == "table_row":
            for chunk in parsed["chunks"]:
                if not is_technical_garbage(chunk):
                    candidates.add(chunk)
                    
        elif parsed["type"] == "text":
            if not is_technical_garbage(parsed["text"]):
                candidates.add(parsed["text"])
                
    return candidates

def get_and_update_glossary(raw_text: str, folder_id: str, api_key: str, github_token: str, github_repo: str) -> dict:
    """Загружает глоссарий из Git, находит новые фразы, переводит только их пачками и пушит в Git."""
    if not github_token or not github_repo:
        st.error("GitHub конфигурация не найдена в Secrets!")
        return {}

    # Устраняем DeprecationWarning для PyGithub
    auth_provider = github.Auth.Token(github_token)
    g = Github(auth=auth_provider)
    repo = g.get_repo(github_repo)
    file_path = "glossary.json"
    
    contents = None
    current_glossary = {}
    
    # 1. Читаем базу данных из текущей ветки
    try:
        contents = repo.get_contents(file_path, ref=TARGET_BRANCH)
        current_glossary = json.loads(contents.decoded_content.decode("utf-8"))
    except json.JSONDecodeError as jde:
        st.error(f"Синтаксическая ошибка в glossary.json на GitHub! Ошибка: {jde}")
        st.stop()
    except Exception as e:
        st.warning(f"Не удалось получить файл глоссария из ветки {TARGET_BRANCH}. Будет создан новый. Ошибка: {e}")
        current_glossary = {}

    # 2. Выделяем всех кандидатов на перевод из документа
    all_candidates = extract_translation_candidates(raw_text)
    # Ищем только те фразы, которых РЕАЛЬНО нет в нашей базе данных
    unknown_candidates = [cand for cand in all_candidates if cand not in current_glossary]
    
    if not unknown_candidates:
        st.success("🎉 Полное совпадение со словарём! В документе нет новых фраз. Расход токенов: 0 рублей.")
        return current_glossary

    st.info(f"🕵️‍♂️ Найдено {len(unknown_candidates)} новых уникальных фраз. Отправляем на экономичный перевод...")
    
    client = openai.OpenAI(
        api_key=api_key,
        base_url="https://ai.api.cloud.yandex.net/v1",
        project=folder_id
    )
    
    # Режем массив незнакомых фраз на небольшие пачки по 30 штук, чтобы JSON не обрывался
    BATCH_SIZE = 30
    new_translations = {}
    
    for i in range(0, len(unknown_candidates), BATCH_SIZE):
        batch = unknown_candidates[i:i+BATCH_SIZE]
        st.caption(f"Перевод пачки {i//BATCH_SIZE + 1} из {len(unknown_candidates)//BATCH_SIZE + 1}...")
        
        batch_json_placeholder = json.dumps({string: "" for string in batch}, ensure_ascii=False, indent=2)
        
        prompt = (
            "Ты — AI-модуль нормализации химической документации (ГОСТ 30333-2022).\n"
            "Переведи предоставленные ключи на русский язык.\n"
            "ТРЕБОВАНИЕ: Верни строго валидный JSON-объект, где ключ — оригинальная английская строка, а значение — русский перевод.\n"
            "Не пиши никаких вступлений или markdown-разметки. Только чистый JSON.\n\n"
            f"Шаблон для заполнения:\n{batch_json_placeholder}"
        )
        
        try:
            response = client.responses.create(
                model=f"gpt://{folder_id}/yandexgpt",
                input=[{"role": "user", "content": prompt}],
                temperature=0.1
            )
            answer = response.output[0].content[0].text.strip()
            answer = re.sub(r'\x60\x60\x60(?:json)?\s*|\s*\x60\x60\x60', '', answer)
            
            parsed_batch = json.loads(answer)
            new_translations.update(parsed_batch)
        except Exception as e:
            st.error(f"Ошибка перевода пачки фраз: {e}")
            continue

    if new_translations:
        # Интегрируем новые переводы в общую базу данных
        current_glossary.update(new_translations)
        updated_content = json.dumps(current_glossary, ensure_ascii=False, indent=4)
        commit_message = f"CAT-Cache обновление: добавлено {len(new_translations)} новых фраз"
        
        try:
            if contents is not None:
                repo.update_file(contents.path, commit_message, updated_content, contents.sha, branch=TARGET_BRANCH)
            else:
                repo.create_file(file_path, commit_message, updated_content, branch=TARGET_BRANCH)
            st.success(f"🧠 База знаний успешно расширена и сохранена в ветку {TARGET_BRANCH}!")
        except Exception as e:
            st.error(f"Не удалось отправить коммит на GitHub: {e}")

    return current_glossary

def assemble_translated_document(text: str, glossary: dict, product_name_ru: str) -> str:
    """
    Финальная высокоточная сборка документа. 
    Берет переводы из кэша, гарантирует совпадение ключей, обрабатывает 
    ручные маркеры склейки/переноса строк и наводит идеальную ГОСТ-верстку.
    """
    cleaned_lines = []
    seen_sections = set()
    stop_words = ['www.', 'http', 'safety data sheet', 'material safety data sheet']
    
    def translate_chunk(chunk: str) -> str:
        if not chunk:
            return ""
        c_clean = chunk.strip()
        if is_technical_garbage(c_clean):
            return c_clean
        return glossary.get(c_clean, c_clean)

    # Флаг-маркер: нужно ли принудительно склеить следующую строку со старой
    merge_next = False

    for line in text.split('\n'):
        line_str = line.strip()
        if not line_str or any(sw in line_str.lower() for sw in stop_words):
            continue
            
        # Защита от мусорных строк, состоящих только из двоеточий или знаков препинания
        if line_str == ":" or re.match(r'^[:\s\-,\|\.]+$', line_str):
            continue

        parsed = parse_line(line_str)
        current_line = ""
        
        # 1. Сборка главных разделов
        if parsed["type"] == "section":
            translated_title = translate_chunk(parsed["text"])
            cleaned_title = translated_title.replace('#', '').strip()
            
            if not cleaned_title.lower().startswith(('раздел', 'section')):
                num_match = re.search(r'\d+', line_str)
                if num_match:
                    cleaned_title = f"РАЗДЕЛ {num_match.group(0)}: {cleaned_title.split(':', 1)[-1].strip()}"
            
            section_marker = " ".join(cleaned_title.split()[:3])
            if section_marker in seen_sections:
                continue
            seen_sections.add(section_marker)
            
            current_line = f"\n# {cleaned_title.upper()}"
            
        # 2. Сборка подразделов (1.1, 2.3.1, 4.2.2 и т.д.) - чистый Markdown без звёздочек
        elif parsed["type"] == "subsection":
            t_key = translate_chunk(parsed["key"])
            num_part = parsed["num"]
            
            if parsed["val"]:
                t_val = translate_chunk(parsed["val"])
                current_line = f"\n## {num_part} {t_key}: {t_val}"
            else:
                current_line = f"\n## {num_part} {t_key}"
                
        # 3. Сборка параметров "Ключ: Значение"
        elif parsed["type"] == "key_value":
            t_key = translate_chunk(parsed["key"])
            t_val = translate_chunk(parsed["val"])
            if t_key and t_val:
                current_line = f"**{t_key}:** {t_val}"
            elif t_key:
                current_line = f"**{t_key}:**"
                
        # 4. Сборка табличных строк в структурированную сетку
        elif parsed["type"] == "table_row":
            t_chunks = [translate_chunk(c) for c in parsed["chunks"]]
            row_str = " | ".join(t_chunks)
            current_line = f"| {row_str} |"
            
        # 5. Сборка обычного текста и списков
        elif parsed["type"] == "text":
            t_text = translate_chunk(parsed["text"])
            prefix = "- " if parsed["is_bullet"] else ""
            current_line = f"{prefix}{t_text}"
            
        if not current_line:
            continue

        # --- ОБРАБОТКА МАРКЕРОВ ГЛОССАРИЯ ---
        # 1. Замена как нативных переносов, так и текстовых "\n" на настоящие переводы строк
        current_line = current_line.replace("\\n", "\n")

        # 2. Поиск знака склейки "<<<" в конце переведённой строки
        next_merge = False
        if current_line.endswith("<<<"):
            next_merge = True
            current_line = current_line[:-3].rstrip()

        # 3. Склеивание с предыдущей строкой при наличии активного флага merge_next
        if merge_next and cleaned_lines:
            last_line = cleaned_lines[-1]
            # Стыкуем аккуратно, убирая лишние концевые пробелы и добавляя ровно один разделительный пробел
            cleaned_lines[-1] = last_line.rstrip() + " " + current_line.lstrip()
        else:
            cleaned_lines.append(current_line)

        # Передаем состояние флага склейки на следующую итерацию цикла
        merge_next = next_merge
            
    final_markdown = '\n'.join(cleaned_lines)
    # Финальный штрих: глобально заменяем имя продукта на его официальное русское имя по ТЗ
    return re.sub(r'ТРИМЕТИЛОЛ\s*ПРОПАН|TRIMETHYLOL\s*PROPANE', product_name_ru, final_markdown, flags=re.IGNORECASE).strip()

def make_formatted_docx(markdown_text: str, product_name_ru: str, product_cas: str):
    """Сборщик Word-документа по ГОСТ-стилистике"""
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.39)
        section.bottom_margin = Inches(0.39)
        section.left_margin = Inches(0.39)
        section.right_margin = Inches(0.39)
        
        hp = section.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run(f"{product_name_ru} | Паспорт безопасности химической продукции")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8.5)
        hrun.font.italic = True
        hrun.font.color.rgb = RGBColor(128, 128, 128)
        
        fp = section.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        frun = fp.add_run("www.spanlab.in                                                                                  ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(128, 128, 128)
        
    p_header1 = doc.add_paragraph()
    p_header1.paragraph_format.space_before = Pt(0)
    p_header1.paragraph_format.space_after = Pt(2)
    run_h1 = p_header1.add_run("Паспорт безопасности материала")
    run_h1.bold = True
    run_h1.font.name = 'Arial'
    run_h1.font.size = Pt(16)
    
    p_header2 = doc.add_paragraph()
    p_header2.paragraph_format.space_before = Pt(0)
    p_header2.paragraph_format.space_after = Pt(2)
    run_h2 = p_header2.add_run(product_name_ru)
    run_h2.bold = True
    run_h2.font.name = 'Arial'
    run_h2.font.size = Pt(16)
    
    p_header3 = doc.add_paragraph()
    p_header3.paragraph_format.space_before = Pt(0)
    p_header3.paragraph_format.space_after = Pt(18)
    run_h3 = p_header3.add_run(f"CAS № {product_cas.strip()}")
    run_h3.bold = True
    run_h3.font.name = 'Arial'
    run_h3.font.size = Pt(16)
    
    pPr = p_header3._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')       
    bottom.set(qn('w:space'), '8')    
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    DARK_BLUE = RGBColor(0, 51, 102)
    lines = markdown_text.split('\n')
    
    for line in lines:
        cleaned_line = line.strip()
        if not cleaned_line:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        
        if cleaned_line.startswith('# '):
            # 1. Удаляем маркер заголовка И полностью вырезаем любые звёздочки
            text_content = cleaned_line.replace('# ', '').replace('**', '').strip()
            run = p.add_run(text_content)
            run.bold = True
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        elif cleaned_line.startswith('## '):
            # 2. Удаляем маркер подраздела И полностью вырезаем любые звёздочки
            text_content = cleaned_line.replace('## ', '').replace('**', '').strip()
            run = p.add_run(text_content)
            run.bold = True  
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.color.rgb = DARK_BLUE
            p.paragraph_format.space_before = Pt(6)
            
        else:
            if cleaned_line.startswith('- '):
                cleaned_line = cleaned_line.replace('- ', '', 1)
                p.paragraph_format.left_indent = Inches(0.25)
            
            # 3. Для обычного текста делим по звёздочкам, чтобы сделать нужные слова жирными
            parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    bold_text = part.replace('**', '')
                    run = p.add_run(bold_text)
                    run.bold = True
                else:
                    # Если звёздочка осталась «одинокой» или сломалась — вырезаем её из текста
                    clean_part = part.replace('**', '')
                    run = p.add_run(clean_part)
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def render_glossary_tab():
    """Вкладка с интерактивной таблицей для управления глоссарием на GitHub"""
    st.header("Управление глоссарием")
    st.caption("Здесь вы можете просматривать, изменять и удалять записи словаря. Изменения автоматически улетят на GitHub.")
    
    if not GITHUB_TOKEN or not GITHUB_REPO:
        st.error("GitHub конфигурации не найдены в Streamlit Secrets!")
        return

    try:
        # Устраняем DeprecationWarning для PyGithub
        auth_provider = github.Auth.Token(GITHUB_TOKEN)
        g = Github(auth=auth_provider)
        repo = g.get_repo(GITHUB_REPO)
        contents = repo.get_contents("glossary.json", ref=TARGET_BRANCH)
        glossary_data = json.loads(contents.decoded_content.decode("utf-8"))
        
        data_list = [{"Оригинал (English)": k, "Перевод (Russian)": v} for k, v in glossary_data.items()]
        data_list.sort(key=lambda x: x["Оригинал (English)"].lower())
        
        # Исправляем use_container_width на width='stretch' по новому стандарту Streamlit
        edited_df = st.data_editor(data_list, width="stretch", num_rows="dynamic")
        
        if st.button("💾 Сохранить изменения в словаре", type="primary"):
            updated_dict = {row["Оригинал (English)"]: row["Перевод (Russian)"] for row in edited_df if row["Оригинал (English)"]}
            new_content = json.dumps(updated_dict, ensure_ascii=False, indent=4)
            
            repo.update_file(
                contents.path, 
                "Ручное редактирование глоссария через интерфейс", 
                new_content, 
                contents.sha,
                branch=TARGET_BRANCH
            )
            st.success(f"Словарь успешно обновлен в ветке {TARGET_BRANCH} на GitHub! Изменения применятся к следующим переводам.")
            
    except Exception as e:
        st.error(f"Не удалось загрузить данные с GitHub. Ошибка: {e}")

# --- Инициализация состояния ---
if "raw_text" not in st.session_state:
    st.session_state.raw_text = ""
if "translated_text" not in st.session_state:
    st.session_state.translated_text = ""
if "current_glossary_cache" not in st.session_state:
    st.session_state.current_glossary_cache = {}
if "file_name_output" not in st.session_state:
    st.session_state.file_name_output = "MSDS_RU_Translated"

def reset_state():
    st.session_state.raw_text = ""
    st.session_state.translated_text = ""
    st.session_state.current_glossary_cache = {}

# --- Основной Интерфейс ---
st.title("🧪 MSDS Translator — Premium AI Studio")

# Разбиваем на две вкладки
tab_main, tab_glossary = st.tabs(["🔄 Переводчик MSDS", "📚 Редактор глоссария"])

with tab_main:
    st.sidebar.header("Параметры продукта")
    product_name_ru = st.sidebar.text_input("Официальное название продукта (RU):", value="ТРИМЕТИЛОЛПРОПАН")
    product_cas = st.sidebar.text_input("Номер CAS:", value="77-99-6")

    # --- ШАГ 1 ---
    st.header("Шаг 1: Загрузка исходного MSDS (EN)")
    input_method = st.radio("Способ загрузки:", ("Загрузить файл (DOCX / PDF / TXT)", "Вставить текст вручную"), on_change=reset_state)

    if input_method == "Вставить текст вручную":
        inserted_text = st.text_area("Вставьте текст MSDS на английском языке:", height=250, placeholder="SECTION 1: Identification...")
        if inserted_text:
            st.session_state.raw_text = inserted_text
    else:
        uploaded_file = st.file_uploader("Выберите файл", type=["docx", "pdf", "txt"])
        if uploaded_file is not None:
            # Отсекаем расширение (берём всё, что до последней точки)
            base_name = uploaded_file.name.rsplit('.', 1)[0]
            st.session_state.file_name_output = f"{base_name}_RU"
            
            if uploaded_file.name.endswith(".txt"):
                st.session_state.raw_text = str(uploaded_file.read(), "utf-8")
            elif uploaded_file.name.endswith(".docx"):
                st.session_state.raw_text = extract_raw_xml_text_from_zip(uploaded_file.read())
            elif uploaded_file.name.endswith(".pdf"):
                with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
                    st.session_state.raw_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

    if st.session_state.raw_text:
        with st.expander("Просмотр извлеченного оригинального текста (Шаг 1)", expanded=False):
            st.text_area("Оригинал без изменений:", value=st.session_state.raw_text, height=200, disabled=True, key="raw_preview")

    st.divider()

    # --- ШАГ 2 ---
    st.header("Шаг 2: Выравнивание и нормализация табличной структуры")
    st.caption("Автоматически сверяется с глоссарием на GitHub, добавляет новые заголовки и выравнивает текст.")

    # Исправляем use_container_width на width='stretch' по новому стандарту Streamlit
    if st.button("🔧 Запустить интеллектуальный анализ", type="secondary", width="stretch"):
        if st.session_state.raw_text:
            with st.spinner("Синхронизация с базой знаний Git и перевод неизвестных фраз..."):
                # Находим новые фразы, шлем в YandexGPT только их, коммитим в указанную ветку
                st.session_state.current_glossary_cache = get_and_update_glossary(
                    st.session_state.raw_text, 
                    FOLDER_ID, 
                    API_KEY, 
                    GITHUB_TOKEN,
                    GITHUB_REPO
                )
            st.success("База знаний обновлена! Все фразы текущего документа теперь есть в кэше.")
        else:
            st.warning("Сначала загрузите или вставьте исходный текст на Шаге 1.")

    st.divider()

    # --- ШАГ 3 ---
    st.header("Шаг 3: Мгновенная сборка перевода из кэша (0 рублей)")
    st.caption("Документ собирается локально без повторных обращений к нейросети.")

    # Исправляем use_container_width на width='stretch' по новому стандарту Streamlit
    if st.button("🚀 Собрать готовый документ", type="primary", width="stretch"):
        # Проверяем, запущен ли кэш в текущей сессии
        if not st.session_state.current_glossary_cache:
            with st.spinner("Загрузка активного словаря..."):
                st.session_state.current_glossary_cache = get_and_update_glossary(
                    st.session_state.raw_text, FOLDER_ID, API_KEY, GITHUB_TOKEN, GITHUB_REPO
                )
        
        if st.session_state.current_glossary_cache:
            with st.spinner("Локальная сборка ГОСТ-структуры..."):
                # Собираем перевод за 1 секунду абсолютно бесплатно!
                st.session_state.translated_text = assemble_translated_document(
                    st.session_state.raw_text, 
                    st.session_state.current_glossary_cache, 
                    product_name_ru
                )
            st.success("Документ успешно собран!")
        else:
            st.warning("Не удалось подготовить кэш для сборки. Выполните Шаг 2.")

    if st.session_state.translated_text:
        with st.expander("Предпросмотр готового перевода Markdown (Шаг 3)", expanded=True):
            st.markdown(st.session_state.translated_text)

    st.divider()

    # --- ШАГ 4 ---
    st.header("Шаг 4: Экспорт в Word")

    if st.session_state.translated_text:
        if "Ошибка" not in st.session_state.translated_text:
            docx_data = make_formatted_docx(st.session_state.translated_text, product_name_ru, product_cas)
            st.download_button(
                label="Скачать отформатированный файл WORD (.docx)",
                data=docx_data,
                file_name=st.session_state.file_name_output if st.session_state.file_name_output.endswith(".docx") else f"{st.session_state.file_name_output}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                width="stretch"
            )
    else:
        st.info("Кнопка скачивания появится здесь, когда Шаг 3 будет успешно выполнен.")

with tab_glossary:
    render_glossary_tab()
